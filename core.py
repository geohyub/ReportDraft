"""Core engine for Processing Report Draft generation."""
import copy
import html
import json
import os
import re
from dataclasses import dataclass, field, asdict
from datetime import datetime
from typing import Optional


@dataclass
class ProcessingStep:
    order: int
    name: str
    description: str = ""
    parameters: dict = field(default_factory=dict)
    stage: str = ""
    rationale: str = ""
    qc_focus: str = ""
    expected_output: str = ""


@dataclass
class ProcessingFlow:
    project_name: str = ""
    client: str = ""
    data_type: str = "SBP"
    vessel: str = ""
    area: str = ""
    date: str = ""
    software: str = "RadExPro"
    software_version: str = ""
    steps: list = field(default_factory=list)
    line_count: int = 0
    notes: str = ""

    @property
    def step_count(self):
        return len(self.steps)


# Common RadExPro processing steps for SBP
DEFAULT_SBP_STEPS = [
    ProcessingStep(1, "Data Input", "SEG-Y data loading and header mapping", {
        "Input format": "SEG-Y Rev 1",
        "Byte order": "Big Endian",
    }),
    ProcessingStep(2, "Geometry Assignment", "Navigation data application and coordinate system setup", {
        "Coordinate system": "WGS84 / UTM",
        "Navigation source": "Embedded in trace headers",
    }),
    ProcessingStep(3, "Trace Editing", "Bad trace removal and data cleanup", {
        "Method": "Manual + automatic spike detection",
    }),
    ProcessingStep(4, "Band-pass Filter", "Frequency filtering to remove noise", {
        "Type": "Ormsby / Butterworth",
        "Low cut": "TBD Hz",
        "High cut": "TBD Hz",
    }),
    ProcessingStep(5, "Gain Application", "Time-varying gain to compensate for amplitude decay", {
        "Type": "AGC / TVG",
        "Window": "TBD ms",
    }),
    ProcessingStep(6, "Swell Filter / Heave Compensation", "Removal of swell-induced noise", {
        "Method": "Swell filter / Static correction",
    }),
    ProcessingStep(7, "Deconvolution", "Source signature removal for improved resolution", {
        "Type": "Predictive / Spiking",
        "Operator length": "TBD ms",
    }),
    ProcessingStep(8, "Migration", "Spatial repositioning of reflectors", {
        "Type": "Stolt / Kirchhoff",
        "Velocity": "TBD m/s",
    }),
    ProcessingStep(9, "Mute / Seafloor Tracking", "Definition of seafloor horizon and muting above", {
        "Method": "Automatic + manual correction",
    }),
    ProcessingStep(10, "SEG-Y Output", "Final processed data export", {
        "Output format": "SEG-Y Rev 1",
        "Sample format": "IEEE 32-bit float",
    }),
]

DEFAULT_UHR_STEPS = [
    ProcessingStep(1, "Data Input", "Multi-channel SEG-D/SEG-Y loading", {
        "Input format": "SEG-Y / SEG-D",
        "Channels": "TBD",
    }),
    ProcessingStep(2, "Geometry Assignment", "Source-receiver geometry and navigation", {
        "Source-receiver offset": "TBD m",
        "Streamer length": "TBD m",
    }),
    ProcessingStep(3, "Trace Editing", "Noisy trace removal and QC", {}),
    ProcessingStep(4, "Band-pass Filter", "Frequency domain filtering", {
        "Low cut": "TBD Hz",
        "High cut": "TBD Hz",
    }),
    ProcessingStep(5, "Gain / AGC", "Amplitude compensation", {
        "AGC window": "TBD ms",
    }),
    ProcessingStep(6, "Deconvolution", "Wavelet compression", {}),
    ProcessingStep(7, "Velocity Analysis", "Velocity picking and NMO correction", {
        "Method": "Semblance analysis",
        "Interval": "Every TBD CDPs",
    }),
    ProcessingStep(8, "NMO Correction", "Normal moveout correction", {}),
    ProcessingStep(9, "CMP Stack", "Common midpoint stacking", {
        "Fold": "TBD",
    }),
    ProcessingStep(10, "Migration", "Post-stack time migration", {
        "Type": "Stolt / Kirchhoff",
        "Velocity model": "Stacking velocities",
    }),
    ProcessingStep(11, "Post-Stack Processing", "Final filtering and scaling", {}),
    ProcessingStep(12, "SEG-Y Output", "Final export", {
        "Output format": "SEG-Y Rev 1",
    }),
]


DEFAULT_MBES_STEPS = [
    ProcessingStep(1, "Data Import", "Raw multibeam data import from acquisition system", {
        "Format": "*.all / *.s7k / *.db / *.kmall",
        "Software": "CARIS HIPS / QPS Qimera",
    }),
    ProcessingStep(2, "Navigation QC", "Vessel position data validation and smoothing", {
        "Position source": "DGPS / PPK",
        "Accuracy": "TBD m",
    }),
    ProcessingStep(3, "Sound Velocity Profile", "SVP application for ray tracing correction", {
        "SVP source": "CTD / SVP cast",
        "Application method": "Nearest in distance/time",
    }),
    ProcessingStep(4, "Tide Correction", "Water level correction to chart datum", {
        "Tide source": "Observed / Predicted / RTK",
        "Datum": "LAT / MSL / CD",
    }),
    ProcessingStep(5, "Vessel Configuration", "Lever arm offsets and mounting angles", {
        "IMU-transducer offsets": "TBD m",
        "Patch test applied": "Yes / No",
    }),
    ProcessingStep(6, "Swath Editing", "Outlier removal and noise cleaning", {
        "Method": "Automatic + Manual",
        "Filter type": "Median / CUBE / Surface",
    }),
    ProcessingStep(7, "Surface Generation", "Bathymetric surface creation", {
        "Resolution": "TBD m",
        "Method": "Weighted Mean / CUBE / Shoal Bias",
    }),
    ProcessingStep(8, "Quality Assessment", "IHO S-44 or project-specific accuracy check", {
        "Standard": "IHO S-44 Special Order / Order 1a / 1b",
        "THU/TVU": "TBD m",
    }),
    ProcessingStep(9, "Export", "Final product export", {
        "Output format": "GeoTIFF / BAG / XYZ / ASCII",
        "Coordinate system": "WGS84 / UTM Zone TBD",
    }),
]


DEFAULT_MAG_STEPS = [
    ProcessingStep(1, "Data Import", "Raw magnetometer data import", {
        "Format": "ASCII / Binary / MagLog",
        "Sensor type": "Proton / Overhauser / Cesium",
    }),
    ProcessingStep(2, "Navigation Merge", "Merge mag readings with vessel positioning", {
        "Layback": "TBD m",
        "Position source": "DGPS",
    }),
    ProcessingStep(3, "Diurnal Correction", "Removal of temporal magnetic field variation", {
        "Base station": "Yes / No",
        "Source": "Local base / INTERMAGNET",
    }),
    ProcessingStep(4, "Spike Removal", "Remove erroneous readings and dropouts", {
        "Method": "Threshold / 4th difference",
        "Threshold": "TBD nT",
    }),
    ProcessingStep(5, "IGRF Removal", "Subtraction of International Geomagnetic Reference Field", {
        "IGRF model": "IGRF-13 / WMM",
        "Epoch": "TBD",
    }),
    ProcessingStep(6, "Heading Correction", "Compensation for sensor heading effects", {
        "Method": "Line-based correction",
    }),
    ProcessingStep(7, "Leveling", "Cross-line/tie-line leveling for consistency", {
        "Method": "Statistical / Polynomial",
        "Tie lines used": "TBD",
    }),
    ProcessingStep(8, "Gridding", "Magnetic anomaly grid generation", {
        "Grid size": "TBD m",
        "Method": "Minimum curvature / Kriging",
    }),
    ProcessingStep(9, "Analytic Signal", "Derivative products for interpretation", {
        "Products": "Total gradient / Tilt derivative / RTP",
    }),
    ProcessingStep(10, "Export", "Final product export", {
        "Output format": "GeoTIFF / XYZ / Geosoft Grid",
    }),
]


DEFAULT_SSS_STEPS = [
    ProcessingStep(1, "Data Import", "Raw side scan sonar data import", {
        "Format": "XTF / JSF / SDF",
        "Frequency": "TBD kHz",
    }),
    ProcessingStep(2, "Navigation QC", "Towfish position and layback correction", {
        "Layback method": "Cable out / USBL",
        "Positioning": "DGPS",
    }),
    ProcessingStep(3, "Slant Range Correction", "Geometric correction for water column", {
        "Method": "Flat bottom / DTM-based",
        "Towfish altitude": "TBD m",
    }),
    ProcessingStep(4, "Gain Normalization", "Along-track and across-track gain correction", {
        "TVG": "Applied / Not applied",
        "EGN": "Empirical Gain Normalization",
    }),
    ProcessingStep(5, "Speed Correction", "Along-track distortion correction", {
        "Speed source": "GPS SOG / vessel speed",
    }),
    ProcessingStep(6, "Bottom Tracking", "Automatic seafloor detection and nadir removal", {
        "Method": "Automatic + Manual QC",
    }),
    ProcessingStep(7, "Mosaicking", "Image mosaic generation", {
        "Resolution": "TBD m/pixel",
        "Blending": "Feathered / Nadir priority",
    }),
    ProcessingStep(8, "Contact Detection", "Target identification and reporting", {
        "Method": "Automatic + Manual",
        "Classification": "Boulder / Debris / Cable / Unknown",
    }),
    ProcessingStep(9, "Export", "Final mosaic and contact report export", {
        "Output format": "GeoTIFF / KMZ / Shapefile",
    }),
]


CORE_DATA_TYPES = ["SBP", "UHR", "MBES", "MAG", "SSS"]

CANONICAL_DATA_TYPE_MAP = {
    "UHRS": "UHR",
    "2DHR": "UHR",
    "MULTIBEAM": "MBES",
    "MAGNETICS": "MAG",
    "SIDESCAN": "SSS",
}

ALIASES_BY_DATA_TYPE = {
    "SBP": [],
    "UHR": ["UHRS", "2DHR"],
    "MBES": ["MULTIBEAM"],
    "MAG": ["MAGNETICS"],
    "SSS": ["SIDESCAN"],
}

DATA_TYPE_BRIEFS = {
    "SBP": {
        "label": "Sub-Bottom Profiler",
        "default_software": "RadExPro",
        "narrative_focus": "preserving shallow reflector continuity while reducing positioning, swell, and random trace noise",
        "story": "This template tells a shallow-imaging story: confirm geometry first, clean traces conservatively, and only then sharpen reflectors for interpretation-ready profiles.",
        "why_template": "SBP reports are convincing when they explain not only which filters were used, but why those choices protected seafloor and shallow sub-seafloor events from over-processing.",
        "deliverables": [
            "Processed SBP SEG-Y profiles ready for interpretation",
            "A step-by-step rationale that explains why shallow-event continuity was preserved",
            "QC notes on navigation, swell/heave handling, and final reflector legibility",
        ],
        "qc_checks": [
            "Navigation and line continuity are still defensible after geometry setup",
            "Noise suppression improves readability without flattening shallow reflectors",
            "The exported profiles match the interpretation objective named in the draft",
        ],
    },
    "UHR": {
        "label": "Ultra-High Resolution Seismic",
        "default_software": "RadExPro",
        "narrative_focus": "building a defensible imaging sequence from geometry through velocity work and stack/migration",
        "story": "This template explains a multi-channel processing story: stabilize geometry, condition traces, justify velocity-driven imaging, and document how the final stack supports interpretation.",
        "why_template": "UHR drafts feel trustworthy when readers can follow the imaging logic from raw gathers to migrated output rather than seeing a disconnected list of signal-processing steps.",
        "deliverables": [
            "Processed stack and migrated seismic sections",
            "A velocity and imaging narrative suitable for internal review or client reporting",
            "QC notes showing how gather conditioning, stack quality, and migration intent connect",
        ],
        "qc_checks": [
            "Geometry and offsets support the later NMO and stack assumptions",
            "Velocity analysis choices are traceable to the reported imaging objective",
            "Post-stack output reads as a coherent imaging workflow rather than a black box",
        ],
    },
    "MBES": {
        "label": "Multibeam Echo Sounder",
        "default_software": "CARIS HIPS and SIPS",
        "narrative_focus": "turning raw swath soundings into a controlled bathymetric surface with explicit correction logic",
        "story": "This template frames MBES processing as a correction-and-surface story: prove the positioning and environmental corrections, clean the swath responsibly, then document how the final grid was accepted.",
        "why_template": "MBES drafts are easier to trust when each correction stage explains what survey bias it addresses before the reader sees the final surface and QA claim.",
        "deliverables": [
            "Cleaned bathymetric surface and export package",
            "A correction log covering navigation, SVP, tide, and vessel configuration decisions",
            "QA wording that links surface generation choices to the final acceptance standard",
        ],
        "qc_checks": [
            "Each applied correction clearly states what positional or depth bias it removes",
            "Swath cleaning is described as a quality decision, not just a software action",
            "The reported surface and export coordinate system align with the stated project objective",
        ],
    },
    "MAG": {
        "label": "Marine Magnetics",
        "default_software": "Oasis Montaj",
        "narrative_focus": "turning raw total-field observations into a leveled anomaly product that can support interpretation",
        "story": "This template tells a correction-first magnetic story: merge navigation confidently, remove temporal and regional field effects, then level and grid the anomaly product for interpretation.",
        "why_template": "MAG drafts feel credible when each correction explains which non-geologic component was removed before the anomaly grid is presented as meaningful.",
        "deliverables": [
            "Corrected and leveled magnetic anomaly dataset",
            "A draft narrative explaining diurnal, IGRF, heading, and leveling decisions",
            "Interpretation-ready grid and derivative products with clear processing provenance",
        ],
        "qc_checks": [
            "Navigation merge and layback handling are explicit before anomaly work begins",
            "Temporal and regional field removals are clearly separated from geologic interpretation",
            "Leveling and gridding are presented as consistency steps, not unexplained finishing moves",
        ],
    },
    "SSS": {
        "label": "Side Scan Sonar",
        "default_software": "SonarWiz / CARIS",
        "narrative_focus": "converting raw sonar records into a georeferenced mosaic and contact package that operators can trust",
        "story": "This template presents SSS processing as an image-quality story: establish positioning, correct geometry and gain, then document how the mosaic and contact outputs were made reliable for review.",
        "why_template": "SSS drafts work best when they explain how geometric and radiometric corrections improved interpretability before presenting the final mosaic and target list.",
        "deliverables": [
            "Processed sonar mosaic and contact package",
            "A draft explanation of layback, slant-range, gain, and mosaicking choices",
            "QC wording that links image clarity to downstream contact interpretation",
        ],
        "qc_checks": [
            "Towfish positioning and layback assumptions are visible before mosaic claims are made",
            "Gain and bottom-tracking choices explain how image consistency was achieved",
            "The final mosaic/contact outputs read as defensible deliverables, not raw software exports",
        ],
    },
}

STEP_STORY_OVERRIDES = {
    "SBP": {
        "Data Input": {
            "aliases": ["SEG-Y Input", "Input"],
            "stage": "Input & Geometry",
            "rationale": "Locks the trace/header contract before any interpretation-sensitive processing begins.",
            "qc_focus": "Verify format, endian, and trace-header mapping before amplitude work.",
            "expected_output": "Imported SBP lines with trusted headers.",
        },
        "Geometry Assignment": {
            "stage": "Input & Geometry",
            "rationale": "Turns traces into defensible survey positions so later seafloor events can be explained spatially.",
            "qc_focus": "Confirm CRS, navigation source, and line continuity.",
            "expected_output": "Positioned profiles ready for editing.",
        },
        "Trace Editing": {
            "stage": "Signal Conditioning",
            "rationale": "Removes acquisition artifacts before filters begin to shape the signal character.",
            "qc_focus": "Check that bad traces are removed without damaging consistent events.",
            "expected_output": "Cleaner trace ensemble with obvious artifacts removed.",
        },
        "Band-pass Filter": {
            "stage": "Signal Conditioning",
            "rationale": "Suppresses out-of-band noise while protecting the reflector bandwidth needed for shallow interpretation.",
            "qc_focus": "Confirm cut frequencies improve readability without clipping real events.",
            "expected_output": "Frequency-balanced traces with improved signal-to-noise ratio.",
        },
        "Gain Application": {
            "aliases": ["Gain", "Apply AGC"],
            "stage": "Signal Conditioning",
            "rationale": "Balances amplitude decay so the draft can explain deeper reflector visibility more confidently.",
            "qc_focus": "Check that gain clarifies the section without exaggerating noise.",
            "expected_output": "Amplitude-balanced section for downstream sharpening.",
        },
        "Swell Filter / Heave Compensation": {
            "stage": "Signal Conditioning",
            "rationale": "Addresses vessel-motion effects that can make shallow events look inconsistent or unstable.",
            "qc_focus": "Ensure motion compensation improves event continuity rather than introducing striping.",
            "expected_output": "Profiles with reduced swell/heave distortion.",
        },
        "Deconvolution": {
            "stage": "Imaging Enhancement",
            "rationale": "Sharpens the source wavelet so reflector boundaries read more clearly in the final draft.",
            "qc_focus": "Confirm the operator improves resolution without ringing or instability.",
            "expected_output": "Higher-resolution traces with cleaner reflector character.",
        },
        "Migration": {
            "stage": "Imaging Enhancement",
            "rationale": "Repositions dipping or diffractions so the report can describe a geologically plausible image rather than acquisition geometry artifacts.",
            "qc_focus": "Check migration velocity and event focusing.",
            "expected_output": "Spatially corrected SBP image ready for review.",
        },
        "Mute / Seafloor Tracking": {
            "stage": "Interpretation Readiness",
            "rationale": "Makes the seafloor reference explicit and removes non-interpretive energy above it.",
            "qc_focus": "Confirm the tracked horizon follows the real seafloor and not noise bursts.",
            "expected_output": "Profiles with a stable seafloor reference and reduced above-bottom clutter.",
        },
        "SEG-Y Output": {
            "aliases": ["Output", "Export"],
            "stage": "Output & Delivery",
            "rationale": "Packages the processed section in an interpretation-ready format with traceability back to the workflow story.",
            "qc_focus": "Verify output format, sample type, and naming are aligned with downstream use.",
            "expected_output": "Final processed SEG-Y deliverable.",
        },
    },
    "UHR": {
        "Data Input": {
            "aliases": ["Import", "Input"],
            "stage": "Input & Geometry",
            "rationale": "Establishes the gather inventory and channel structure that later imaging depends on.",
            "qc_focus": "Confirm input format, channel count, and trace organization.",
            "expected_output": "Loaded UHR gathers ready for geometry setup.",
        },
        "Geometry Assignment": {
            "stage": "Input & Geometry",
            "rationale": "Builds the source-receiver context needed for velocity analysis, NMO, and stack.",
            "qc_focus": "Check offsets, streamer assumptions, and navigation consistency.",
            "expected_output": "Positioned gathers with defensible acquisition geometry.",
        },
        "Trace Editing": {
            "stage": "Signal Conditioning",
            "rationale": "Removes noisy or dead traces before they contaminate stack energy and velocity picks.",
            "qc_focus": "Confirm edits target artifacts without collapsing useful coverage.",
            "expected_output": "Cleaner gathers for spectral conditioning.",
        },
        "Band-pass Filter": {
            "stage": "Signal Conditioning",
            "rationale": "Shapes the spectrum so later velocity and stack decisions are based on a clearer signal band.",
            "qc_focus": "Check that filtering improves coherence without removing usable bandwidth.",
            "expected_output": "Conditioned gathers with reduced spectral noise.",
        },
        "Gain / AGC": {
            "aliases": ["Gain", "Apply AGC"],
            "stage": "Signal Conditioning",
            "rationale": "Balances energy for gather review and velocity picking.",
            "qc_focus": "Ensure the gain window supports interpretation instead of boosting random noise.",
            "expected_output": "Amplitude-balanced gathers.",
        },
        "Deconvolution": {
            "stage": "Signal Conditioning",
            "rationale": "Compresses the wavelet so moveout and reflector continuity are easier to judge.",
            "qc_focus": "Watch for ringing or unstable phase behavior after deconvolution.",
            "expected_output": "Sharper gathers with improved temporal resolution.",
        },
        "Velocity Analysis": {
            "stage": "Velocity & Imaging",
            "rationale": "Provides the imaging logic that later justifies NMO, stack, and migration choices.",
            "qc_focus": "Document pick spacing and whether semblance events are stable enough to trust.",
            "expected_output": "Velocity picks or model assumptions for the draft.",
        },
        "NMO Correction": {
            "stage": "Velocity & Imaging",
            "rationale": "Applies the selected velocity model so gathers can stack coherently.",
            "qc_focus": "Check residual moveout and stretch behavior after correction.",
            "expected_output": "NMO-corrected gathers ready for stacking.",
        },
        "CMP Stack": {
            "stage": "Velocity & Imaging",
            "rationale": "Turns gather energy into a section that can support structural interpretation.",
            "qc_focus": "Confirm stack continuity and whether noise reduction matches expectations.",
            "expected_output": "Stacked seismic section.",
        },
        "Migration": {
            "stage": "Velocity & Imaging",
            "rationale": "Positions events more realistically so the report can speak about imaging quality, not just stack quality.",
            "qc_focus": "Review event focusing, dip response, and migration-velocity consistency.",
            "expected_output": "Migrated seismic image for interpretation.",
        },
        "Post-Stack Processing": {
            "stage": "Interpretation Readiness",
            "rationale": "Provides final balancing so the migrated section reads clearly in the report and deliverable package.",
            "qc_focus": "Check that post-stack shaping clarifies the image without hiding uncertainty.",
            "expected_output": "Presentation-ready stacked or migrated section.",
        },
        "SEG-Y Output": {
            "aliases": ["Output", "Export"],
            "stage": "Output & Delivery",
            "rationale": "Packages the final imaging result in a reusable format with a workflow the reader can audit.",
            "qc_focus": "Confirm export naming, sample format, and final product selection.",
            "expected_output": "Final UHR SEG-Y deliverable.",
        },
    },
    "MBES": {
        "Data Import": {
            "aliases": ["Import raw .all files", "Import", "Load raw files"],
            "stage": "Input & Corrections",
            "rationale": "Defines the raw sounding inventory and source system before correction logic is applied.",
            "qc_focus": "Confirm file format, survey date grouping, and import completeness.",
            "expected_output": "Imported swath dataset ready for correction.",
        },
        "Navigation QC": {
            "stage": "Input & Corrections",
            "rationale": "Stabilizes vessel positioning so later depth corrections are not undermined by poor horizontal control.",
            "qc_focus": "Check position smoothing, outliers, and reference source quality.",
            "expected_output": "Navigation solution suitable for bathymetric processing.",
        },
        "Sound Velocity Profile": {
            "aliases": ["Apply SVP correction", "SVP correction"],
            "stage": "Input & Corrections",
            "rationale": "Explains how refraction was handled before sounding geometry is judged.",
            "qc_focus": "Confirm SVP source, selection logic, and whether casts are representative.",
            "expected_output": "Soundings corrected for water-column refraction.",
        },
        "Tide Correction": {
            "aliases": ["Apply RTK tides", "Tide correction"],
            "stage": "Input & Corrections",
            "rationale": "Brings depths to a stated datum so the final surface can be compared and delivered consistently.",
            "qc_focus": "Check datum choice, tide source, and timing alignment.",
            "expected_output": "Depths referenced to the target datum.",
        },
        "Vessel Configuration": {
            "stage": "Input & Corrections",
            "rationale": "Makes offsets and mounting geometry explicit so the draft can defend platform corrections.",
            "qc_focus": "Verify offsets, alignment values, and patch-test status.",
            "expected_output": "Survey vessel model aligned with the sonar geometry.",
        },
        "Swath Editing": {
            "aliases": ["Swath editing", "Noise cleaning"],
            "stage": "Cleaning & Surface Build",
            "rationale": "Removes obvious outliers before the surface starts to tell the wrong story about seabed morphology.",
            "qc_focus": "Check that cleaning removes artifacts without carving away valid seabed edges.",
            "expected_output": "Cleaned soundings ready for gridding.",
        },
        "Surface Generation": {
            "aliases": ["Surface generation", "Create surface"],
            "stage": "Cleaning & Surface Build",
            "rationale": "Transforms cleaned soundings into the surface product that the report ultimately defends.",
            "qc_focus": "Confirm grid resolution, method, and density are appropriate for the survey objective.",
            "expected_output": "Bathymetric surface or grid product.",
        },
        "Quality Assessment": {
            "stage": "Acceptance & Delivery",
            "rationale": "Shows how the final surface was judged against a named quality standard instead of being accepted implicitly.",
            "qc_focus": "State which acceptance standard was used and what evidence supports it.",
            "expected_output": "Quality verdict and acceptance notes for the draft.",
        },
        "Export": {
            "aliases": ["Export to BAG", "Export"],
            "stage": "Acceptance & Delivery",
            "rationale": "Packages the accepted surface into the coordinate system and file types expected downstream.",
            "qc_focus": "Check product list, CRS, and delivery naming.",
            "expected_output": "Final MBES export package.",
        },
    },
    "MAG": {
        "Data Import": {
            "aliases": ["Import", "Input"],
            "stage": "Input & Positioning",
            "rationale": "Defines the magnetic dataset and sensor context before any anomaly correction is claimed.",
            "qc_focus": "Confirm file completeness and sensor metadata.",
            "expected_output": "Imported total-field observations.",
        },
        "Navigation Merge": {
            "stage": "Input & Positioning",
            "rationale": "Links field readings to defendable locations so anomaly products can later be interpreted spatially.",
            "qc_focus": "Verify position source, timing alignment, and layback assumptions.",
            "expected_output": "Positioned magnetic records.",
        },
        "Diurnal Correction": {
            "stage": "Field Corrections",
            "rationale": "Removes temporal magnetic variation before the draft speaks about geologic signal.",
            "qc_focus": "Check base-station source and timing consistency.",
            "expected_output": "Time-corrected magnetic observations.",
        },
        "Spike Removal": {
            "stage": "Field Corrections",
            "rationale": "Clears dropouts and obvious noise spikes so later anomaly products are not driven by acquisition artifacts.",
            "qc_focus": "Confirm despiking is selective and traceable.",
            "expected_output": "Cleaned magnetic line data.",
        },
        "IGRF Removal": {
            "stage": "Field Corrections",
            "rationale": "Separates the regional reference field from the local anomaly story the report wants to tell.",
            "qc_focus": "State the chosen IGRF model and epoch.",
            "expected_output": "Residual anomaly values after regional-field removal.",
        },
        "Heading Correction": {
            "stage": "Field Corrections",
            "rationale": "Compensates platform heading effects so line-to-line comparison is defensible.",
            "qc_focus": "Check whether heading bias is reduced consistently across headings.",
            "expected_output": "Heading-corrected magnetic lines.",
        },
        "Leveling": {
            "stage": "Consistency Build",
            "rationale": "Reduces line mismatch so the final grid reads as one survey product rather than separate traverses.",
            "qc_focus": "Review tie-line behavior and the residual mismatch after leveling.",
            "expected_output": "Internally consistent line set for gridding.",
        },
        "Gridding": {
            "stage": "Consistency Build",
            "rationale": "Creates the map product that the report uses to discuss anomalies and interpretation confidence.",
            "qc_focus": "Confirm grid cell size and interpolation suit line spacing.",
            "expected_output": "Magnetic anomaly grid.",
        },
        "Analytic Signal": {
            "stage": "Interpretation Readiness",
            "rationale": "Adds derivative products that make anomaly edges and targets easier to describe.",
            "qc_focus": "Explain which derivative products support the stated interpretation goal.",
            "expected_output": "Secondary interpretation products such as analytic signal or tilt derivative.",
        },
        "Export": {
            "stage": "Output & Delivery",
            "rationale": "Packages corrected grids and anomaly products for mapping, interpretation, or sharing.",
            "qc_focus": "Check final format, naming, and metadata completeness.",
            "expected_output": "Final MAG delivery package.",
        },
    },
    "SSS": {
        "Data Import": {
            "aliases": ["Import", "Input"],
            "stage": "Input & Positioning",
            "rationale": "Defines the raw image inventory before any geometric or radiometric correction is claimed.",
            "qc_focus": "Confirm file completeness, format, and sonar-frequency context.",
            "expected_output": "Imported side-scan records ready for correction.",
        },
        "Navigation QC": {
            "stage": "Input & Positioning",
            "rationale": "Makes towfish position assumptions explicit so the mosaic can later be trusted spatially.",
            "qc_focus": "Check layback logic, navigation source, and positioning consistency.",
            "expected_output": "Positioned side-scan lines with documented towfish handling.",
        },
        "Slant Range Correction": {
            "stage": "Geometric Conditioning",
            "rationale": "Removes water-column geometry distortion so the seabed image can be read more naturally.",
            "qc_focus": "Verify bottom pick quality and whether nadir handling is stable.",
            "expected_output": "Geometrically corrected seabed imagery.",
        },
        "Gain Normalization": {
            "stage": "Radiometric Conditioning",
            "rationale": "Balances brightness so the draft can explain image consistency rather than raw acquisition variability.",
            "qc_focus": "Check that normalization improves comparability without flattening useful contrast.",
            "expected_output": "Radiometrically balanced imagery.",
        },
        "Speed Correction": {
            "stage": "Geometric Conditioning",
            "rationale": "Removes along-track distortion before mosaic scale and target shape are discussed.",
            "qc_focus": "Confirm corrected pixel aspect follows vessel motion realistically.",
            "expected_output": "Speed-stabilized sonar line images.",
        },
        "Bottom Tracking": {
            "stage": "Radiometric Conditioning",
            "rationale": "Defines the seafloor and nadir region so later mosaicking uses the right image footprint.",
            "qc_focus": "Check tracking stability and nadir removal behavior.",
            "expected_output": "Tracked seabed extent for each line.",
        },
        "Mosaicking": {
            "stage": "Mosaic Build",
            "rationale": "Builds the view the reader actually sees in the report and uses for contact review.",
            "qc_focus": "Confirm overlap handling, blending, and final mosaic resolution.",
            "expected_output": "Survey mosaic ready for interpretation or delivery.",
        },
        "Contact Detection": {
            "stage": "Interpretation Readiness",
            "rationale": "Turns image interpretation into a structured target list the report can discuss.",
            "qc_focus": "Check that contact logic and classification rules are explicit.",
            "expected_output": "Target/contact register linked to the mosaic.",
        },
        "Export": {
            "stage": "Output & Delivery",
            "rationale": "Packages the mosaic and contact set in formats that support client review and downstream GIS use.",
            "qc_focus": "Verify output formats, georeferencing, and file naming.",
            "expected_output": "Final SSS export package.",
        },
    },
}


# Supported data types and their templates
DATA_TYPE_TEMPLATES = {
    "SBP": DEFAULT_SBP_STEPS,
    "UHR": DEFAULT_UHR_STEPS,
    "UHRS": DEFAULT_UHR_STEPS,
    "2DHR": DEFAULT_UHR_STEPS,
    "MBES": DEFAULT_MBES_STEPS,
    "MULTIBEAM": DEFAULT_MBES_STEPS,
    "MAG": DEFAULT_MAG_STEPS,
    "MAGNETICS": DEFAULT_MAG_STEPS,
    "SSS": DEFAULT_SSS_STEPS,
    "SIDESCAN": DEFAULT_SSS_STEPS,
}

SUPPORTED_DATA_TYPES = sorted(set(DATA_TYPE_TEMPLATES.keys()))


def canonicalize_data_type(data_type):
    """Return the canonical base data type used by guided drafting."""
    normalized = (data_type or "SBP").strip().upper()
    return CANONICAL_DATA_TYPE_MAP.get(normalized, normalized)


def _normalize_name(value):
    return re.sub(r"[^a-z0-9]+", "", (value or "").lower())


def _tokenize_name(value):
    return {
        token
        for token in re.findall(r"[a-z0-9]+", (value or "").lower())
        if token not in {"the", "and", "for", "to", "of", "apply"}
    }


def _is_tbd_value(value):
    text = str(value or "").strip().lower()
    if not text:
        return True
    return any(marker in text for marker in ("tbd", "tbc", "to be confirmed", "to be decided"))


def _get_step_story(data_type, step_name):
    canonical = canonicalize_data_type(data_type)
    stories = STEP_STORY_OVERRIDES.get(canonical, {})
    normalized_target = _normalize_name(step_name)
    if not normalized_target:
        return {}

    best_story = {}
    best_score = 0.0
    target_tokens = _tokenize_name(step_name)

    for story_name, story in stories.items():
        candidates = [story_name] + list(story.get("aliases", []))
        for candidate in candidates:
            if _normalize_name(candidate) == normalized_target:
                return story

            candidate_tokens = _tokenize_name(candidate)
            if not candidate_tokens or not target_tokens:
                continue

            overlap = len(candidate_tokens & target_tokens)
            score = overlap / max(len(candidate_tokens), len(target_tokens))
            if score > best_score:
                best_score = score
                best_story = story

    return best_story if best_score >= 0.5 else {}


def enrich_flow(flow):
    """Populate guided-drafting metadata for each step."""
    canonical = canonicalize_data_type(flow.data_type)
    flow.data_type = canonical
    for step in flow.steps:
        story = _get_step_story(canonical, step.name)
        if story:
            if not step.stage:
                step.stage = story.get("stage", "")
            if not step.rationale:
                step.rationale = story.get("rationale", "")
            if not step.qc_focus:
                step.qc_focus = story.get("qc_focus", "")
            if not step.expected_output:
                step.expected_output = story.get("expected_output", "")

    if not flow.software:
        flow.software = DATA_TYPE_BRIEFS.get(canonical, {}).get("default_software", "RadExPro")
    return flow


def parse_processing_log(log_text):
    """Parse a text-based processing log into ProcessingFlow.

    Supports formats:
    - Simple numbered steps: "1. Step Name - Description"
    - Key: Value pairs for parameters
    - YAML-like structures
    """
    flow = ProcessingFlow()
    steps = []
    current_step = None
    order = 0

    for line in log_text.split("\n"):
        line = line.strip()
        if not line:
            continue

        # Check for project metadata
        lower = line.lower()
        if current_step is None:
            if lower.startswith("project:"):
                flow.project_name = line.split(":", 1)[1].strip()
                continue
            elif lower.startswith("client:"):
                flow.client = line.split(":", 1)[1].strip()
                continue
            elif lower.startswith("data type:") or lower.startswith("type:"):
                flow.data_type = line.split(":", 1)[1].strip()
                continue
            elif lower.startswith("vessel:"):
                flow.vessel = line.split(":", 1)[1].strip()
                continue
            elif lower.startswith("area:"):
                flow.area = line.split(":", 1)[1].strip()
                continue
            elif lower.startswith("software:"):
                flow.software = line.split(":", 1)[1].strip()
                continue
            elif lower.startswith("lines:") or lower.startswith("line count:"):
                try:
                    flow.line_count = int(re.search(r"\d+", line.split(":", 1)[1]).group())
                except (AttributeError, ValueError):
                    pass
                continue

        # Check for numbered step
        step_match = re.match(r"^(\d+)[.)\s]+(.+)", line)
        if step_match:
            if current_step:
                steps.append(current_step)
            order += 1
            name = step_match.group(2).strip()
            desc = ""
            if " - " in name:
                name, desc = name.split(" - ", 1)
            elif ": " in name:
                name, desc = name.split(": ", 1)
            current_step = ProcessingStep(order=order, name=name.strip(), description=desc.strip())
            continue

        # Check for parameter (indented key:value)
        if current_step and ":" in line and not line.startswith("#"):
            key, _, val = line.partition(":")
            key = key.strip()
            val = val.strip()
            if key and val:
                current_step.parameters[key] = val

    if current_step:
        steps.append(current_step)

    flow.steps = steps
    return enrich_flow(flow)


def generate_docx_report(flow, output_path):
    """Generate a Word document processing report draft."""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    flow = enrich_flow(copy.deepcopy(flow))
    context = build_flow_context(flow)

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.name = "Calibri"
    styles["Heading 3"].font.size = Pt(11)

    # Title
    title = doc.add_heading(level=0)
    title.add_run("Processing Report Draft")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        f"{flow.project_name or context['label']} | {flow.data_type} | {flow.software or context['label']}"
    )
    subtitle_run.font.size = Pt(12)

    readiness = doc.add_paragraph()
    readiness.alignment = WD_ALIGN_PARAGRAPH.CENTER
    readiness_run = readiness.add_run(context["readiness"]["label"])
    readiness_run.bold = True

    readiness_detail = doc.add_paragraph()
    readiness_detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
    readiness_detail.add_run(context["readiness"]["detail"])

    doc.add_paragraph()

    # Project Overview
    doc.add_heading("1. Project Overview", level=1)
    info_data = [
        ("Project Name", flow.project_name or "TBD"),
        ("Client", flow.client or "TBD"),
        ("Data Type", f"{flow.data_type or 'TBD'} ({context['label']})"),
        ("Vessel", flow.vessel or "TBD"),
        ("Survey Area", flow.area or "TBD"),
        ("Processing Software", f"{flow.software} {flow.software_version}".strip()),
        ("Number of Lines", str(flow.line_count) if flow.line_count else "TBD"),
        ("Report Date", datetime.now().strftime("%Y-%m-%d")),
        ("Draft Readiness", context["readiness"]["label"]),
    ]

    table = doc.add_table(rows=len(info_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (key, val) in enumerate(info_data):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = val
        # Bold the key
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.add_paragraph()

    doc.add_heading("2. Draft Briefing", level=1)
    for section_info in context["report_sections"]:
        doc.add_heading(section_info["title"], level=2)
        doc.add_paragraph(section_info["body"])

    doc.add_heading("3. Workflow Story by Stage", level=1)
    if not flow.steps:
        doc.add_paragraph("No processing steps have been defined yet.")
    else:
        for group in context["stage_groups"]:
            doc.add_heading(group["stage"], level=2)
            doc.add_paragraph(group["summary"])

            for step in group["steps"]:
                doc.add_heading(f"Step {step['order']}. {step['name']}", level=3)
                step_obj = next((item for item in flow.steps if item.order == step["order"]), None)
                if step_obj and step_obj.description:
                    doc.add_paragraph(step_obj.description)

                story_rows = []
                if step_obj and step_obj.rationale:
                    story_rows.append(("Why this step is used", step_obj.rationale))
                if step_obj and step_obj.qc_focus:
                    story_rows.append(("What to check", step_obj.qc_focus))
                if step_obj and step_obj.expected_output:
                    story_rows.append(("What this stage should produce", step_obj.expected_output))

                if story_rows:
                    story_table = doc.add_table(rows=len(story_rows), cols=2)
                    story_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for row_idx, (label, value) in enumerate(story_rows):
                        story_table.rows[row_idx].cells[0].text = label
                        story_table.rows[row_idx].cells[1].text = value
                        for paragraph in story_table.rows[row_idx].cells[0].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

                if step_obj and step_obj.parameters:
                    param_table = doc.add_table(rows=len(step_obj.parameters) + 1, cols=2)
                    param_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    param_table.rows[0].cells[0].text = "Parameter"
                    param_table.rows[0].cells[1].text = "Value"
                    for cell in param_table.rows[0].cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

                    for param_idx, (key, value) in enumerate(step_obj.parameters.items(), 1):
                        param_table.rows[param_idx].cells[0].text = key
                        param_table.rows[param_idx].cells[1].text = str(value)

                doc.add_paragraph()

    doc.add_heading("4. Deliverables and Review Prompts", level=1)
    for deliverable in context["deliverables"]:
        doc.add_paragraph(deliverable, style="List Bullet")

    if context["open_items"]:
        doc.add_heading("Review Prompts", level=2)
        for item in context["open_items"]:
            doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("5. Appendix - Step Summary", level=1)
    summary_table = doc.add_table(rows=max(len(flow.steps), 1) + 1, cols=4)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_headers = ["Order", "Stage", "Step", "QC Focus"]
    for idx, header in enumerate(summary_headers):
        summary_table.rows[0].cells[idx].text = header
        for paragraph in summary_table.rows[0].cells[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    if flow.steps:
        for row_idx, step in enumerate(flow.steps, 1):
            summary_table.rows[row_idx].cells[0].text = str(step.order)
            summary_table.rows[row_idx].cells[1].text = step.stage or "-"
            summary_table.rows[row_idx].cells[2].text = step.name
            summary_table.rows[row_idx].cells[3].text = step.qc_focus or "-"
    else:
        summary_table.rows[1].cells[0].text = "-"
        summary_table.rows[1].cells[1].text = "-"
        summary_table.rows[1].cells[2].text = "No steps defined"
        summary_table.rows[1].cells[3].text = "-"

    doc.save(output_path)
    return output_path


def generate_flow_from_template(data_type="SBP"):
    """Get a default processing flow template."""
    dt_upper = canonicalize_data_type(data_type)
    flow = ProcessingFlow(data_type=dt_upper)

    template = DATA_TYPE_TEMPLATES.get(dt_upper, DEFAULT_SBP_STEPS)
    flow.steps = [copy.deepcopy(s) for s in template]

    # Set default software based on data type
    if dt_upper in ("MBES", "MULTIBEAM"):
        flow.software = "CARIS HIPS and SIPS"
    elif dt_upper in ("MAG", "MAGNETICS"):
        flow.software = "Oasis Montaj"
    elif dt_upper in ("SSS", "SIDESCAN"):
        flow.software = "SonarWiz / CARIS"
    else:
        flow.software = "RadExPro"

    return enrich_flow(flow)


def generate_text_report(flow):
    """Generate a plain text processing report."""
    flow = enrich_flow(copy.deepcopy(flow))
    context = build_flow_context(flow)
    lines = []
    lines.append("=" * 70)
    lines.append("  ProcessingReportDraft - Data Processing Report Draft")
    lines.append("=" * 70)
    lines.append("")

    # Metadata
    lines.append(f"  Project:    {flow.project_name or 'TBD'}")
    lines.append(f"  Client:     {flow.client or 'TBD'}")
    lines.append(f"  Data Type:  {flow.data_type or 'TBD'}")
    lines.append(f"  Vessel:     {flow.vessel or 'TBD'}")
    lines.append(f"  Area:       {flow.area or 'TBD'}")
    lines.append(f"  Software:   {flow.software} {flow.software_version}".rstrip())
    lines.append(f"  Lines:      {flow.line_count or 'TBD'}")
    lines.append(f"  Status:     {context['readiness']['label']}")
    lines.append("")
    lines.append("-" * 70)
    lines.append("  Draft Briefing")
    lines.append("-" * 70)
    for section_info in context["report_sections"]:
        lines.append(f"  {section_info['title']}:")
        lines.append(f"    {section_info['body']}")
        lines.append("")

    lines.append("-" * 70)
    lines.append(f"  Workflow by Stage ({flow.step_count} steps)")
    lines.append("-" * 70)

    for group in context["stage_groups"]:
        lines.append(f"  Stage: {group['stage']}")
        lines.append(f"    {group['summary']}")
        for step in group["steps"]:
            lines.append(f"    Step {step['order']}: {step['name']}")
            if step["description"]:
                lines.append(f"      Description: {step['description']}")
            if step["rationale"]:
                lines.append(f"      Why: {step['rationale']}")
            if step["qc_focus"]:
                lines.append(f"      QC Focus: {step['qc_focus']}")
            if step["expected_output"]:
                lines.append(f"      Expected Output: {step['expected_output']}")
            if step_obj := next((item for item in flow.steps if item.order == step["order"]), None):
                for key, value in step_obj.parameters.items():
                    lines.append(f"      - {key}: {value}")
        lines.append("")

    if context["deliverables"]:
        lines.append("-" * 70)
        lines.append("  Deliverables")
        lines.append("-" * 70)
        for item in context["deliverables"]:
            lines.append(f"  - {item}")
        lines.append("")

    if context["open_items"]:
        lines.append("-" * 70)
        lines.append("  Open Items")
        lines.append("-" * 70)
        for item in context["open_items"]:
            lines.append(f"  - {item}")

    lines.append("")
    lines.append("=" * 70)
    if flow.notes:
        lines.append(f"\n  Notes: {flow.notes}")
    return "\n".join(lines)


def get_supported_types():
    """Return list of supported data types with info."""
    info = {}
    for dt in CORE_DATA_TYPES:
        flow = generate_flow_from_template(dt)
        brief = DATA_TYPE_BRIEFS.get(dt, {})
        stages = []
        for step in flow.steps:
            if step.stage and step.stage not in stages:
                stages.append(step.stage)
        info[dt] = {
            "name": dt,
            "label": brief.get("label", dt),
            "default_software": brief.get("default_software", flow.software),
            "step_count": len(flow.steps),
            "steps_preview": [s.name for s in flow.steps[:5]],
            "story": brief.get("story", ""),
            "why_template": brief.get("why_template", ""),
            "narrative_focus": brief.get("narrative_focus", ""),
            "deliverables": brief.get("deliverables", []),
            "aliases": ALIASES_BY_DATA_TYPE.get(dt, []),
            "stages": stages,
        }
    return info


# ── Custom Step Editing ──

def _renumber_steps(flow):
    """Re-number all steps sequentially starting from 1."""
    for i, step in enumerate(flow.steps):
        step.order = i + 1
    return flow


def add_custom_step(
    flow,
    name,
    description="",
    parameters=None,
    position=None,
    stage="",
    rationale="",
    qc_focus="",
    expected_output="",
):
    """Add a custom step to flow. If position is None, append to end.
    Auto-sets order numbers. Returns updated flow."""
    if parameters is None:
        parameters = {}
    new_step = ProcessingStep(
        order=0,
        name=name,
        description=description,
        parameters=parameters,
        stage=stage,
        rationale=rationale,
        qc_focus=qc_focus,
        expected_output=expected_output,
    )
    if position is None or position > len(flow.steps):
        flow.steps.append(new_step)
    else:
        idx = max(0, position - 1)
        flow.steps.insert(idx, new_step)
    _renumber_steps(flow)
    return enrich_flow(flow)


def remove_step(flow, order):
    """Remove step by order number and re-number remaining steps.
    Returns updated flow."""
    flow.steps = [s for s in flow.steps if s.order != order]
    _renumber_steps(flow)
    return enrich_flow(flow)


def reorder_steps(flow, new_order):
    """Reorder steps. new_order is list of current order numbers in desired sequence.
    Example: [3,1,2] puts step 3 first, then 1, then 2.
    Returns updated flow."""
    order_map = {s.order: s for s in flow.steps}
    # Validate that all order numbers exist
    for o in new_order:
        if o not in order_map:
            raise ValueError(f"Step with order {o} does not exist")
    if len(new_order) != len(flow.steps):
        raise ValueError("new_order must contain all step order numbers")
    flow.steps = [order_map[o] for o in new_order]
    _renumber_steps(flow)
    return enrich_flow(flow)


def update_step(
    flow,
    order,
    name=None,
    description=None,
    parameters=None,
    stage=None,
    rationale=None,
    qc_focus=None,
    expected_output=None,
):
    """Update a specific step's properties. Returns updated flow."""
    for step in flow.steps:
        if step.order == order:
            if name is not None:
                step.name = name
            if description is not None:
                step.description = description
            if parameters is not None:
                step.parameters = parameters
            if stage is not None:
                step.stage = stage
            if rationale is not None:
                step.rationale = rationale
            if qc_focus is not None:
                step.qc_focus = qc_focus
            if expected_output is not None:
                step.expected_output = expected_output
            return enrich_flow(flow)
    raise ValueError(f"Step with order {order} not found")


# ── Flow Comparison ──

@dataclass
class FlowDiff:
    added_steps: list = field(default_factory=list)     # steps in flow2 not in flow1
    removed_steps: list = field(default_factory=list)    # steps in flow1 not in flow2
    modified_steps: list = field(default_factory=list)   # steps with different parameters
    metadata_changes: dict = field(default_factory=dict) # changes to project_name, client, etc.


def compare_flows(flow1, flow2):
    """Compare two ProcessingFlows and return differences."""
    diff = FlowDiff()

    # Compare metadata
    metadata_fields = [
        "project_name", "client", "data_type", "vessel",
        "area", "date", "software", "software_version",
        "line_count", "notes",
    ]
    for f in metadata_fields:
        v1 = getattr(flow1, f)
        v2 = getattr(flow2, f)
        if v1 != v2:
            diff.metadata_changes[f] = {"old": v1, "new": v2}

    # Build name-based maps for step comparison
    names1 = {s.name: s for s in flow1.steps}
    names2 = {s.name: s for s in flow2.steps}

    # Added: in flow2 but not flow1
    for name, step in names2.items():
        if name not in names1:
            diff.added_steps.append(step)

    # Removed: in flow1 but not flow2
    for name, step in names1.items():
        if name not in names2:
            diff.removed_steps.append(step)

    # Modified: same name but different parameters or description
    for name in names1:
        if name in names2:
            s1 = names1[name]
            s2 = names2[name]
            if (
                s1.parameters != s2.parameters
                or s1.description != s2.description
                or s1.stage != s2.stage
                or s1.rationale != s2.rationale
                or s1.qc_focus != s2.qc_focus
                or s1.expected_output != s2.expected_output
            ):
                diff.modified_steps.append({
                    "name": name,
                    "old": {
                        "description": s1.description,
                        "parameters": s1.parameters,
                        "stage": s1.stage,
                        "rationale": s1.rationale,
                        "qc_focus": s1.qc_focus,
                        "expected_output": s1.expected_output,
                    },
                    "new": {
                        "description": s2.description,
                        "parameters": s2.parameters,
                        "stage": s2.stage,
                        "rationale": s2.rationale,
                        "qc_focus": s2.qc_focus,
                        "expected_output": s2.expected_output,
                    },
                })

    return diff


# ── Revision Tracking ──

def _serialize_flow(flow):
    """Serialize a ProcessingFlow to a dict for storage."""
    return {
        "project_name": flow.project_name,
        "client": flow.client,
        "data_type": flow.data_type,
        "vessel": flow.vessel,
        "area": flow.area,
        "date": flow.date,
        "software": flow.software,
        "software_version": flow.software_version,
        "line_count": flow.line_count,
        "notes": flow.notes,
        "steps": [
            {
                "order": s.order,
                "name": s.name,
                "description": s.description,
                "parameters": dict(s.parameters),
                "stage": s.stage,
                "rationale": s.rationale,
                "qc_focus": s.qc_focus,
                "expected_output": s.expected_output,
            }
            for s in flow.steps
        ],
    }


def _deserialize_flow(d):
    """Deserialize a dict back to a ProcessingFlow."""
    flow = ProcessingFlow(
        project_name=d.get("project_name", ""),
        client=d.get("client", ""),
        data_type=d.get("data_type", "SBP"),
        vessel=d.get("vessel", ""),
        area=d.get("area", ""),
        date=d.get("date", ""),
        software=d.get("software", "RadExPro"),
        software_version=d.get("software_version", ""),
        line_count=d.get("line_count", 0),
        notes=d.get("notes", ""),
    )
    flow.steps = [
        ProcessingStep(
            order=s.get("order", i + 1),
            name=s.get("name", ""),
            description=s.get("description", ""),
            parameters=s.get("parameters", {}),
            stage=s.get("stage", ""),
            rationale=s.get("rationale", ""),
            qc_focus=s.get("qc_focus", ""),
            expected_output=s.get("expected_output", ""),
        )
        for i, s in enumerate(d.get("steps", []))
    ]
    return enrich_flow(flow)


@dataclass
class FlowRevision:
    version: int
    timestamp: str
    author: str
    changes: str  # description of changes
    flow_snapshot: dict = field(default_factory=dict)  # serialized flow at this point


class RevisionTracker:
    def __init__(self):
        self.revisions = []  # list of FlowRevision

    def save_revision(self, flow, author="", changes=""):
        """Save current flow as a new revision."""
        version = len(self.revisions) + 1
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        snapshot = _serialize_flow(flow)
        revision = FlowRevision(
            version=version,
            timestamp=timestamp,
            author=author,
            changes=changes,
            flow_snapshot=snapshot,
        )
        self.revisions.append(revision)
        return revision

    def get_revision(self, version):
        """Get a specific revision."""
        for rev in self.revisions:
            if rev.version == version:
                return rev
        raise ValueError(f"Revision version {version} not found")

    def get_history(self):
        """Get list of all revisions with metadata."""
        return [
            {
                "version": r.version,
                "timestamp": r.timestamp,
                "author": r.author,
                "changes": r.changes,
            }
            for r in self.revisions
        ]

    def diff_revisions(self, v1, v2):
        """Get diff between two revision versions."""
        rev1 = self.get_revision(v1)
        rev2 = self.get_revision(v2)
        flow1 = _deserialize_flow(rev1.flow_snapshot)
        flow2 = _deserialize_flow(rev2.flow_snapshot)
        return compare_flows(flow1, flow2)


# ── Step Parameter Validation ──

PARAMETER_VALIDATION_RULES = {
    "SBP": {
        "Band-pass Filter": {
            "Low cut": {"min": 0.1, "max": 500, "type": "float"},
            "High cut": {"min": 100, "max": 20000, "type": "float"},
            "Type": {"allowed": ["Butterworth", "Ormsby", "Zero-phase"], "type": "choice"},
        },
        "Gain Application": {
            "Window": {"min": 10, "max": 5000, "type": "float"},
        },
        "Migration": {
            "Velocity": {"min": 1400, "max": 2000, "type": "float"},
        },
    },
    "UHR": {
        "Band-pass Filter": {
            "Low cut": {"min": 1, "max": 1000, "type": "float"},
            "High cut": {"min": 50, "max": 50000, "type": "float"},
        },
        "Gain / AGC": {
            "AGC window": {"min": 10, "max": 5000, "type": "float"},
        },
        "Velocity Analysis": {
            "Interval": {"min": 1, "max": 500, "type": "int"},
        },
    },
    "MBES": {
        "Navigation QC": {
            "Accuracy": {"min": 0.01, "max": 25, "type": "float"},
        },
        "Surface Generation": {
            "Resolution": {"min": 0.1, "max": 50, "type": "float"},
        },
    },
    "MAG": {
        "Spike Removal": {
            "Threshold": {"min": 1, "max": 10000, "type": "float"},
        },
        "Gridding": {
            "Grid size": {"min": 1, "max": 1000, "type": "float"},
        },
    },
    "SSS": {
        "Slant Range Correction": {
            "Towfish altitude": {"min": 0.1, "max": 500, "type": "float"},
        },
        "Mosaicking": {
            "Resolution": {"min": 0.05, "max": 10, "type": "float"},
        },
    },
}


def _extract_numeric_value(value):
    if isinstance(value, (int, float)):
        return float(value)

    text = str(value or "").replace(",", "")
    match = re.search(r"-?\d+(?:\.\d+)?", text)
    if not match:
        raise ValueError(f"No numeric content found in {value!r}")
    return float(match.group())


def _issue_message(issue):
    return f"{issue['step']} - {issue['parameter']}: {issue['expected']}"


def validate_flow_parameters(flow):
    """Validate all step parameters against known rules.

    Returns: {
        "total_params_checked": N,
        "valid": N,
        "invalid": N,
        "unknown": N,  # params not in rules (not an error)
        "issues": [
            {"step": str, "parameter": str, "value": str, "expected": str, "severity": "error"/"warning"}
        ],
        "score": float  # 0-100, percentage of valid params
    }
    """
    flow = enrich_flow(flow)
    rules = PARAMETER_VALIDATION_RULES.get(canonicalize_data_type(flow.data_type), {})
    issues = []
    total = 0
    valid = 0
    invalid = 0
    unknown = 0
    tbd_parameters = 0

    for step in flow.steps:
        step_rules = {
            _normalize_name(rule_name): rule
            for rule_name, rule in rules.get(step.name, {}).items()
        }
        for param_name, param_value in step.parameters.items():
            total += 1

            if _is_tbd_value(param_value):
                tbd_parameters += 1
                issues.append({
                    "step": step.name,
                    "parameter": param_name,
                    "value": str(param_value),
                    "expected": "Replace placeholder with project-specific value",
                    "severity": "warning",
                })
                continue

            normalized_param_name = _normalize_name(param_name)
            rule = step_rules.get(normalized_param_name)
            if rule is None:
                for rule_name, candidate_rule in step_rules.items():
                    if rule_name in normalized_param_name or normalized_param_name in rule_name:
                        rule = candidate_rule
                        break
            if rule is None:
                unknown += 1
                continue
            try:
                if rule["type"] == "choice":
                    normalized_value = str(param_value).lower()
                    if not any(allowed.lower() in normalized_value for allowed in rule["allowed"]):
                        invalid += 1
                        issues.append({
                            "step": step.name,
                            "parameter": param_name,
                            "value": str(param_value),
                            "expected": f"One of: {rule['allowed']}",
                            "severity": "error",
                        })
                    else:
                        valid += 1
                elif rule["type"] in ("float", "int"):
                    num_val = _extract_numeric_value(param_value)
                    if num_val < rule["min"] or num_val > rule["max"]:
                        invalid += 1
                        issues.append({
                            "step": step.name,
                            "parameter": param_name,
                            "value": str(param_value),
                            "expected": f"{rule['min']} - {rule['max']}",
                            "severity": "warning",
                        })
                    else:
                        valid += 1
                else:
                    unknown += 1
            except (ValueError, TypeError):
                invalid += 1
                issues.append({
                    "step": step.name,
                    "parameter": param_name,
                    "value": str(param_value),
                    "expected": f"Numeric ({rule['type']})",
                    "severity": "error",
                })

    checked = valid + invalid
    score = (valid / checked * 100) if checked > 0 else 100.0
    errors = [_issue_message(issue) for issue in issues if issue["severity"] == "error"]
    warnings = [_issue_message(issue) for issue in issues if issue["severity"] == "warning"]
    return {
        "total_params_checked": total,
        "valid": valid,
        "invalid": invalid,
        "unknown": unknown,
        "tbd_parameters": tbd_parameters,
        "issues": issues,
        "errors": errors,
        "warnings": warnings,
        "score": round(score, 1),
    }


# ── Flow Statistics & Analytics ──

def get_flow_statistics(flow):
    """Compute comprehensive statistics about a processing flow.

    Returns: {
        "step_count": N,
        "total_parameters": N,
        "avg_params_per_step": float,
        "steps_with_descriptions": N,
        "steps_without_descriptions": N,
        "completeness_score": float,  # 0-100 based on descriptions and params present
        "data_type": str,
        "software": str,
        "has_metadata": bool,  # project, client, vessel all filled
        "parameter_types": {"filter": N, "velocity": N, "output": N, ...},  # categorize params
    }
    """
    flow = enrich_flow(flow)
    step_count = len(flow.steps)
    total_parameters = sum(len(s.parameters) for s in flow.steps)
    avg_params = (total_parameters / step_count) if step_count > 0 else 0.0
    steps_with_desc = sum(1 for s in flow.steps if s.description)
    steps_without_desc = step_count - steps_with_desc
    steps_with_guidance = sum(1 for s in flow.steps if s.rationale or s.qc_focus or s.expected_output)
    tbd_parameters = sum(1 for s in flow.steps for value in s.parameters.values() if _is_tbd_value(value))
    filled_parameters = max(total_parameters - tbd_parameters, 0)
    stage_count = len({s.stage for s in flow.steps if s.stage})

    # Completeness score: based on descriptions and parameter counts
    if step_count == 0:
        completeness_score = 0.0
    else:
        desc_ratio = steps_with_desc / step_count
        param_ratio = min(avg_params / 2.0, 1.0)  # normalize: 2+ params per step = 100%
        guidance_ratio = steps_with_guidance / step_count
        completeness_score = round((desc_ratio * 35 + param_ratio * 35 + guidance_ratio * 30), 1)

    metadata_labels = {
        "project_name": "Project name",
        "client": "Client",
        "vessel": "Vessel",
        "area": "Survey area",
    }
    missing_metadata = [
        label for field_name, label in metadata_labels.items() if not getattr(flow, field_name)
    ]
    if not flow.line_count:
        missing_metadata.append("Line count")

    metadata_completion = round(
        ((len(metadata_labels) + 1 - len(missing_metadata)) / (len(metadata_labels) + 1)) * 100,
        1,
    )
    parameter_completion = round(
        (filled_parameters / total_parameters * 100) if total_parameters > 0 else 100.0,
        1,
    )
    readiness_score = round(
        completeness_score * 0.35 + parameter_completion * 0.40 + metadata_completion * 0.25,
        1,
    )
    has_metadata = bool(flow.project_name and flow.client and flow.vessel)

    # Categorize parameters by keyword
    param_categories = {
        "filter": 0, "velocity": 0, "output": 0, "input": 0,
        "correction": 0, "other": 0,
    }
    filter_keywords = ["filter", "freq", "cut", "band", "pass"]
    velocity_keywords = ["velocity", "nmo", "moveout", "cdp"]
    output_keywords = ["output", "export", "format", "sample"]
    input_keywords = ["input", "import", "load", "byte", "channel"]
    correction_keywords = ["correction", "tide", "svp", "heave", "diurnal", "igrf"]

    for step in flow.steps:
        for param_name in step.parameters:
            pn_lower = param_name.lower()
            categorized = False
            for kw in filter_keywords:
                if kw in pn_lower:
                    param_categories["filter"] += 1
                    categorized = True
                    break
            if not categorized:
                for kw in velocity_keywords:
                    if kw in pn_lower:
                        param_categories["velocity"] += 1
                        categorized = True
                        break
            if not categorized:
                for kw in output_keywords:
                    if kw in pn_lower:
                        param_categories["output"] += 1
                        categorized = True
                        break
            if not categorized:
                for kw in input_keywords:
                    if kw in pn_lower:
                        param_categories["input"] += 1
                        categorized = True
                        break
            if not categorized:
                for kw in correction_keywords:
                    if kw in pn_lower:
                        param_categories["correction"] += 1
                        categorized = True
                        break
            if not categorized:
                param_categories["other"] += 1

    return {
        "step_count": step_count,
        "total_parameters": total_parameters,
        "avg_params_per_step": round(avg_params, 2),
        "steps_with_descriptions": steps_with_desc,
        "steps_without_descriptions": steps_without_desc,
        "steps_with_guidance": steps_with_guidance,
        "completeness_score": completeness_score,
        "data_type": flow.data_type,
        "software": flow.software,
        "has_metadata": has_metadata,
        "missing_metadata": missing_metadata,
        "metadata_completion": metadata_completion,
        "parameter_completion": parameter_completion,
        "draft_readiness": readiness_score,
        "tbd_parameters": tbd_parameters,
        "filled_parameters": filled_parameters,
        "stage_count": stage_count,
        "parameter_types": param_categories,
    }


def compare_flow_statistics(flows):
    """Compare statistics across multiple flows.

    Input: list of ProcessingFlow objects
    Returns: {
        "flow_count": N,
        "comparison": [
            {"data_type": str, "step_count": N, "total_parameters": N, "completeness_score": float}
        ],
        "avg_step_count": float,
        "avg_completeness": float,
        "most_detailed": str,  # data_type with highest completeness
        "least_detailed": str,
    }
    """
    comparison = []
    for flow in flows:
        stats = get_flow_statistics(flow)
        comparison.append({
            "data_type": stats["data_type"],
            "step_count": stats["step_count"],
            "total_parameters": stats["total_parameters"],
            "completeness_score": stats["completeness_score"],
        })

    flow_count = len(flows)
    avg_step_count = round(
        sum(c["step_count"] for c in comparison) / flow_count, 2
    ) if flow_count > 0 else 0.0
    avg_completeness = round(
        sum(c["completeness_score"] for c in comparison) / flow_count, 2
    ) if flow_count > 0 else 0.0

    most_detailed = ""
    least_detailed = ""
    if comparison:
        sorted_by_score = sorted(comparison, key=lambda c: c["completeness_score"])
        most_detailed = sorted_by_score[-1]["data_type"]
        least_detailed = sorted_by_score[0]["data_type"]

    return {
        "flow_count": flow_count,
        "comparison": comparison,
        "avg_step_count": avg_step_count,
        "avg_completeness": avg_completeness,
        "most_detailed": most_detailed,
        "least_detailed": least_detailed,
    }


def _build_stage_groups(flow):
    groups = []
    by_stage = {}

    for step in flow.steps:
        stage = step.stage or "Custom Workflow"
        if stage not in by_stage:
            by_stage[stage] = {
                "stage": stage,
                "summary": "",
                "step_count": 0,
                "steps": [],
            }
            groups.append(by_stage[stage])

        group = by_stage[stage]
        group["steps"].append({
            "order": step.order,
            "name": step.name,
            "description": step.description,
            "rationale": step.rationale,
            "qc_focus": step.qc_focus,
            "expected_output": step.expected_output,
        })
        group["step_count"] += 1

    for group in groups:
        narratives = [item["rationale"] for item in group["steps"] if item["rationale"]]
        outputs = [item["expected_output"] for item in group["steps"] if item["expected_output"]]
        if narratives:
            group["summary"] = narratives[0]
        elif outputs:
            group["summary"] = outputs[0]
        else:
            group["summary"] = f"{group['stage']} work captured in the current draft."

    return groups


def build_report_sections(flow, context=None):
    flow = enrich_flow(flow)
    if context is None:
        context = build_flow_context(flow)

    return [
        {
            "title": "Executive Summary",
            "body": context["executive_summary"],
        },
        {
            "title": "Processing Strategy",
            "body": context["processing_strategy"],
        },
        {
            "title": "QC and Review Focus",
            "body": context["qc_story"],
        },
        {
            "title": "Open Draft Items",
            "body": context["open_story"],
        },
    ]


def build_flow_context(flow):
    flow = enrich_flow(flow)
    canonical = canonicalize_data_type(flow.data_type)
    brief = DATA_TYPE_BRIEFS.get(canonical, {
        "label": canonical or "Processing",
        "story": "Build a transparent processing draft that explains the sequence and the intended output.",
        "why_template": "Readers should understand what problem each step solves before they read the final deliverable summary.",
        "narrative_focus": "connecting workflow decisions to the final output",
        "deliverables": [],
        "qc_checks": [],
    })
    stats = get_flow_statistics(flow)
    validation = validate_flow_parameters(flow)
    stage_groups = _build_stage_groups(flow)

    project_ref = flow.project_name or "this project"
    line_text = f" across {flow.line_count} lines" if flow.line_count else ""
    executive_summary = (
        f"This {brief['label']} processing draft for {project_ref}{line_text} is organized around "
        f"{brief['narrative_focus']}. The workflow currently captures {flow.step_count} steps grouped into "
        f"{len(stage_groups)} stages using {flow.software or brief.get('default_software', 'the stated software')}."
    )
    processing_strategy = f"{brief['story']} {brief['why_template']}"

    qc_points = brief.get("qc_checks", [])
    if qc_points:
        qc_story = "Key review points: " + "; ".join(qc_points)
    else:
        qc_story = "Key review points should explain how each major correction or imaging step was verified."

    open_items = []
    for label in stats["missing_metadata"]:
        open_items.append(f"Confirm {label.lower()} for the final issue draft.")
    for issue in validation["issues"]:
        prefix = "Resolve" if issue["severity"] == "error" else "Confirm"
        item = f"{prefix} {issue['step']} - {issue['parameter']}: {issue['expected']}."
        if item not in open_items:
            open_items.append(item)

    if not open_items:
        open_story = "No blocking placeholders were detected in the current draft. Final reviewer wording and project-specific values can be refined from this baseline."
    else:
        open_story = "Open items before issue: " + " ".join(open_items[:6])

    if validation["invalid"] > 0 or len(stats["missing_metadata"]) >= 3 or stats["tbd_parameters"] >= max(4, flow.step_count // 2 or 1):
        readiness = {
            "label": "Guided draft - key project values still needed",
            "tone": "warn",
            "detail": "The workflow logic is visible, but the draft still needs project-specific metadata or technical values before it reads like a finished report.",
            "score": stats["draft_readiness"],
        }
    elif stats["missing_metadata"] or stats["tbd_parameters"]:
        readiness = {
            "label": "Structured draft - final technical fill-in still needed",
            "tone": "info",
            "detail": "The report story is coherent, and the remaining work is mostly replacing placeholders and confirming project metadata.",
            "score": stats["draft_readiness"],
        }
    else:
        readiness = {
            "label": "Report-ready draft skeleton",
            "tone": "ok",
            "detail": "The workflow is documented with guidance, outputs, and QC emphasis. Final reviewer polish can happen directly in the draft.",
            "score": stats["draft_readiness"],
        }

    context = {
        "canonical_data_type": canonical,
        "label": brief.get("label", canonical),
        "headline": f"{canonical} workflow story focused on {brief['narrative_focus']}",
        "summary": brief.get("story", ""),
        "why_template": brief.get("why_template", ""),
        "narrative_focus": brief.get("narrative_focus", ""),
        "deliverables": brief.get("deliverables", []),
        "qc_checks": brief.get("qc_checks", []),
        "stage_groups": stage_groups,
        "open_items": open_items,
        "readiness": readiness,
        "executive_summary": executive_summary,
        "processing_strategy": processing_strategy,
        "qc_story": qc_story,
        "open_story": open_story,
        "statistics": stats,
        "validation": validation,
    }
    context["report_sections"] = [
        {"title": "Executive Summary", "body": executive_summary},
        {"title": "Processing Strategy", "body": processing_strategy},
        {"title": "QC and Review Focus", "body": qc_story},
        {"title": "Open Draft Items", "body": open_story},
    ]
    return context


# ── Export to Multiple Formats ──

def generate_excel_report(flow, output_path):
    """Generate Excel workbook from a processing flow.

    Sheet 1 'Overview': Project info table (metadata)
    Sheet 2 'Processing Steps': order, name, description, all parameters as columns
    Sheet 3 'Validation': parameter validation results
    Returns output path.
    """
    import openpyxl
    from openpyxl.styles import Font, PatternFill

    flow = enrich_flow(copy.deepcopy(flow))
    context = build_flow_context(flow)
    validation = context["validation"]
    stats = context["statistics"]

    wb = openpyxl.Workbook()

    # ── Sheet 1: Overview ──
    ws1 = wb.active
    ws1.title = "Overview"
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font_white = Font(bold=True, size=12, color="FFFFFF")

    ws1.append(["Field", "Value"])
    for cell in ws1[1]:
        cell.font = header_font_white
        cell.fill = header_fill

    info_rows = [
        ("Project Name", flow.project_name or "TBD"),
        ("Client", flow.client or "TBD"),
        ("Data Type", f"{flow.data_type} ({context['label']})"),
        ("Vessel", flow.vessel or "TBD"),
        ("Survey Area", flow.area or "TBD"),
        ("Software", f"{flow.software} {flow.software_version}".strip()),
        ("Number of Lines", str(flow.line_count) if flow.line_count else "TBD"),
        ("Report Date", datetime.now().strftime("%Y-%m-%d")),
        ("Total Steps", str(len(flow.steps))),
        ("Draft Status", context["readiness"]["label"]),
        ("Draft Readiness Score", f"{stats['draft_readiness']}%"),
        ("Placeholder Parameters", stats["tbd_parameters"]),
    ]
    for key, val in info_rows:
        ws1.append([key, val])

    ws1.column_dimensions["A"].width = 24
    ws1.column_dimensions["B"].width = 60

    # ── Sheet 2: Draft Story ──
    ws_story = wb.create_sheet("Draft Story")
    ws_story.append(["Section", "Narrative"])
    for cell in ws_story[1]:
        cell.font = header_font_white
        cell.fill = header_fill

    for section_info in context["report_sections"]:
        ws_story.append([section_info["title"], section_info["body"]])

    ws_story.append([])
    ws_story.append(["Recommended Deliverables", ""])
    deliverable_header_row = ws_story.max_row
    for cell in ws_story[deliverable_header_row]:
        cell.font = header_font_white
        cell.fill = header_fill
    for item in context["deliverables"]:
        ws_story.append(["Deliverable", item])

    if context["open_items"]:
        ws_story.append([])
        ws_story.append(["Open Draft Items", ""])
        open_header_row = ws_story.max_row
        for cell in ws_story[open_header_row]:
            cell.font = header_font_white
            cell.fill = header_fill
        for item in context["open_items"]:
            ws_story.append(["Action", item])

    ws_story.column_dimensions["A"].width = 24
    ws_story.column_dimensions["B"].width = 90

    # ── Sheet 3: Processing Steps ──
    ws2 = wb.create_sheet("Processing Steps")

    all_param_names = []
    for step in flow.steps:
        for param_name in step.parameters:
            if param_name not in all_param_names:
                all_param_names.append(param_name)

    headers = [
        "Order", "Stage", "Name", "Description",
        "Why Used", "QC Focus", "Expected Output",
    ] + all_param_names
    ws2.append(headers)
    for cell in ws2[1]:
        cell.font = header_font_white
        cell.fill = header_fill

    for step in flow.steps:
        row = [
            step.order,
            step.stage,
            step.name,
            step.description,
            step.rationale,
            step.qc_focus,
            step.expected_output,
        ]
        for param_name in all_param_names:
            row.append(str(step.parameters.get(param_name, "")))
        ws2.append(row)

    for col in ("A", "B", "C", "D", "E", "F", "G"):
        ws2.column_dimensions[col].width = 24

    # ── Sheet 4: Validation ──
    ws3 = wb.create_sheet("Validation")
    ws3.append(["Validation Summary"])
    ws3["A1"].font = Font(bold=True, size=14)
    ws3.append(["Total Parameters Checked", validation["total_params_checked"]])
    ws3.append(["Valid", validation["valid"]])
    ws3.append(["Invalid", validation["invalid"]])
    ws3.append(["Unknown", validation["unknown"]])
    ws3.append(["Placeholder Parameters", validation["tbd_parameters"]])
    ws3.append(["Score", f"{validation['score']}%"])
    ws3.append([])

    if validation["issues"]:
        ws3.append(["Step", "Parameter", "Value", "Expected", "Severity"])
        for cell in ws3[ws3.max_row]:
            cell.font = header_font_white
            cell.fill = header_fill
        for issue in validation["issues"]:
            ws3.append([
                issue["step"],
                issue["parameter"],
                issue["value"],
                issue["expected"],
                issue["severity"],
            ])

    ws3.column_dimensions["A"].width = 28
    ws3.column_dimensions["B"].width = 22
    ws3.column_dimensions["C"].width = 18
    ws3.column_dimensions["D"].width = 48
    ws3.column_dimensions["E"].width = 14

    wb.save(output_path)
    return output_path


def generate_json_export(flow):
    """Export flow as formatted JSON string.

    Returns: JSON string with proper formatting.
    """
    flow = enrich_flow(copy.deepcopy(flow))
    context = build_flow_context(flow)
    data = {
        "project_name": flow.project_name,
        "client": flow.client,
        "data_type": flow.data_type,
        "vessel": flow.vessel,
        "area": flow.area,
        "date": flow.date,
        "software": flow.software,
        "software_version": flow.software_version,
        "line_count": flow.line_count,
        "notes": flow.notes,
        "step_count": len(flow.steps),
        "steps": [
            {
                "order": s.order,
                "name": s.name,
                "description": s.description,
                "parameters": dict(s.parameters),
                "stage": s.stage,
                "rationale": s.rationale,
                "qc_focus": s.qc_focus,
                "expected_output": s.expected_output,
            }
            for s in flow.steps
        ],
        "statistics": context["statistics"],
        "validation": context["validation"],
        "context": context,
    }
    return json.dumps(data, indent=2, ensure_ascii=False)


def generate_html_report(flow):
    """Generate standalone HTML report (for email/sharing).

    Returns: HTML string with embedded CSS styling.
    Contains: project info table, step cards, parameter tables.
    """
    flow = enrich_flow(copy.deepcopy(flow))
    context = build_flow_context(flow)
    esc = html.escape
    readiness_bg = {
        "warn": "var(--warn-soft)",
        "info": "var(--accent-soft)",
        "ok": "var(--ok-soft)",
    }.get(context["readiness"]["tone"], "var(--accent-soft)")
    readiness_fg = {
        "warn": "var(--warn)",
        "info": "var(--accent)",
        "ok": "var(--ok)",
    }.get(context["readiness"]["tone"], "var(--accent)")

    section_html = "".join(
        f"""
        <div class="narrative-card">
            <h3>{esc(section_info['title'])}</h3>
            <p>{esc(section_info['body'])}</p>
        </div>
        """
        for section_info in context["report_sections"]
    )

    deliverable_html = "".join(f"<li>{esc(item)}</li>" for item in context["deliverables"])
    open_items_html = "".join(f"<li>{esc(item)}</li>" for item in context["open_items"])

    stage_html = ""
    for group in context["stage_groups"]:
        step_cards = ""
        for step in group["steps"]:
            params_html = ""
            step_obj = next((item for item in flow.steps if item.order == step["order"]), None)
            if step_obj and step_obj.parameters:
                param_rows = "".join(
                    f"<tr><td>{esc(key)}</td><td>{esc(str(value))}</td></tr>"
                    for key, value in step_obj.parameters.items()
                )
                params_html = f"""
                <table class="param-table">
                    <tr><th>Parameter</th><th>Value</th></tr>
                    {param_rows}
                </table>
                """

            step_cards += f"""
            <div class="step-card">
                <div class="step-header">
                    <span class="step-order">Step {step['order']}</span>
                    <h4>{esc(step['name'])}</h4>
                </div>
                <p class="step-desc">{esc(step['description'] or '')}</p>
                <div class="step-meta"><strong>Why:</strong> {esc(step['rationale'] or 'Add project-specific rationale.')}</div>
                <div class="step-meta"><strong>QC Focus:</strong> {esc(step['qc_focus'] or 'Add review focus.')}</div>
                <div class="step-meta"><strong>Expected Output:</strong> {esc(step['expected_output'] or 'Add expected output.')}</div>
                {params_html}
            </div>
            """

        stage_html += f"""
        <section class="stage-block">
            <div class="stage-header">
                <div>
                    <span class="stage-kicker">{esc(group['stage'])}</span>
                    <h3>{esc(group['stage'])}</h3>
                </div>
                <span class="stage-count">{group['step_count']} steps</span>
            </div>
            <p class="stage-summary">{esc(group['summary'])}</p>
            {step_cards}
        </section>
        """

    html_output = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Processing Report Draft - {esc(flow.data_type or '')}</title>
    <style>
        :root {{
            --ink: #162033;
            --muted: #5b6b82;
            --line: #d9e1ec;
            --soft: #f4f7fb;
            --accent: #1f4e78;
            --accent-soft: #e8f0f8;
            --warn: #b45309;
            --warn-soft: #fff7ed;
            --ok: #166534;
            --ok-soft: #ecfdf3;
        }}
        * {{ box-sizing: border-box; }}
        body {{
            margin: 0;
            background: #edf2f7;
            color: var(--ink);
            font-family: Calibri, 'Segoe UI', Arial, sans-serif;
            line-height: 1.55;
        }}
        .page {{
            max-width: 1040px;
            margin: 28px auto;
            background: #fff;
            padding: 36px 42px 48px;
            box-shadow: 0 16px 40px rgba(15, 23, 42, 0.08);
        }}
        .hero {{
            border-bottom: 2px solid var(--line);
            padding-bottom: 20px;
            margin-bottom: 28px;
        }}
        .hero h1 {{
            margin: 0 0 8px;
            font-size: 30px;
            color: var(--accent);
        }}
        .hero p {{
            margin: 0;
            color: var(--muted);
        }}
        .status-pill {{
            display: inline-block;
            margin-top: 12px;
            padding: 8px 12px;
            border-radius: 999px;
            background: {readiness_bg};
            color: {readiness_fg};
            font-weight: 700;
            font-size: 14px;
        }}
        h2 {{
            margin: 32px 0 14px;
            font-size: 22px;
            color: var(--accent);
        }}
        h3 {{
            margin: 0 0 8px;
            font-size: 18px;
            color: var(--ink);
        }}
        .overview-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(220px, 1fr));
            gap: 12px;
        }}
        .overview-card, .narrative-card, .stage-block, .sidebar-card {{
            border: 1px solid var(--line);
            background: var(--soft);
            padding: 16px 18px;
            border-radius: 14px;
        }}
        .overview-card .label {{
            font-size: 12px;
            text-transform: uppercase;
            color: var(--muted);
            letter-spacing: 0.08em;
        }}
        .overview-card .value {{
            margin-top: 4px;
            font-size: 16px;
            font-weight: 700;
        }}
        .story-grid {{
            display: grid;
            grid-template-columns: 2fr 1fr;
            gap: 16px;
            align-items: start;
        }}
        .story-stack {{
            display: grid;
            gap: 14px;
        }}
        .sidebar-card ul {{
            margin: 12px 0 0;
            padding-left: 18px;
        }}
        .sidebar-card li {{
            margin-bottom: 8px;
        }}
        .stage-block {{
            margin-bottom: 18px;
            background: #fff;
        }}
        .stage-header {{
            display: flex;
            justify-content: space-between;
            gap: 12px;
            align-items: baseline;
            margin-bottom: 10px;
        }}
        .stage-kicker {{
            display: inline-block;
            font-size: 12px;
            text-transform: uppercase;
            letter-spacing: 0.08em;
            color: var(--muted);
            margin-bottom: 4px;
        }}
        .stage-count {{
            font-size: 13px;
            font-weight: 700;
            color: var(--accent);
        }}
        .stage-summary {{
            margin: 0 0 14px;
            color: var(--muted);
        }}
        .step-card {{
            border: 1px solid var(--line);
            border-left: 5px solid var(--accent);
            border-radius: 12px;
            padding: 14px 16px;
            margin-bottom: 12px;
            background: #fff;
        }}
        .step-header {{
            display: flex;
            align-items: center;
            gap: 10px;
            margin-bottom: 6px;
        }}
        .step-order {{
            display: inline-flex;
            align-items: center;
            justify-content: center;
            min-width: 76px;
            padding: 4px 10px;
            border-radius: 999px;
            background: var(--accent-soft);
            color: var(--accent);
            font-size: 12px;
            font-weight: 700;
        }}
        .step-header h4 {{
            margin: 0;
            font-size: 17px;
        }}
        .step-desc {{
            margin: 0 0 10px;
            color: var(--muted);
        }}
        .step-meta {{
            margin-bottom: 6px;
            font-size: 14px;
        }}
        .param-table {{
            width: 100%;
            border-collapse: collapse;
            margin-top: 12px;
        }}
        .param-table th, .param-table td {{
            border: 1px solid var(--line);
            padding: 8px 10px;
            text-align: left;
            vertical-align: top;
        }}
        .param-table th {{
            background: var(--accent);
            color: #fff;
            font-size: 13px;
        }}
        .footer {{
            margin-top: 28px;
            padding-top: 16px;
            border-top: 1px solid var(--line);
            color: var(--muted);
            font-size: 13px;
            text-align: center;
        }}
        @media (max-width: 900px) {{
            .page {{ padding: 24px; }}
            .story-grid {{ grid-template-columns: 1fr; }}
        }}
    </style>
</head>
<body>
    <div class="page">
        <section class="hero">
            <h1>Processing Report Draft</h1>
            <p>{esc(context['executive_summary'])}</p>
            <div class="status-pill">{esc(context['readiness']['label'])}</div>
        </section>

        <h2>Project Overview</h2>
        <div class="overview-grid">
            <div class="overview-card"><div class="label">Project</div><div class="value">{esc(flow.project_name or 'TBD')}</div></div>
            <div class="overview-card"><div class="label">Client</div><div class="value">{esc(flow.client or 'TBD')}</div></div>
            <div class="overview-card"><div class="label">Data Type</div><div class="value">{esc(flow.data_type)} ({esc(context['label'])})</div></div>
            <div class="overview-card"><div class="label">Software</div><div class="value">{esc(flow.software or 'TBD')}</div></div>
            <div class="overview-card"><div class="label">Lines</div><div class="value">{esc(str(flow.line_count) if flow.line_count else 'TBD')}</div></div>
            <div class="overview-card"><div class="label">Placeholders</div><div class="value">{context['statistics']['tbd_parameters']}</div></div>
        </div>

        <h2>Draft Story</h2>
        <div class="story-grid">
            <div class="story-stack">
                {section_html}
            </div>
            <div class="story-stack">
                <div class="sidebar-card">
                    <h3>Recommended Deliverables</h3>
                    <ul>{deliverable_html or '<li>Add project-specific deliverables.</li>'}</ul>
                </div>
                <div class="sidebar-card">
                    <h3>Open Draft Items</h3>
                    <ul>{open_items_html or '<li>No blocking open items detected.</li>'}</ul>
                </div>
            </div>
        </div>

        <h2>Workflow Story by Stage</h2>
        {stage_html}

        <div class="footer">
            Generated by ProcessingReportDraft on {datetime.now().strftime('%Y-%m-%d')}
        </div>
    </div>
</body>
</html>"""
    return html_output


# ── Bulk Template Operations ──

def generate_all_templates(project_name="", client="", vessel="", area=""):
    """Generate templates for ALL data types at once with common metadata.

    Returns: {
        "SBP": ProcessingFlow,
        "UHR": ProcessingFlow,
        "MBES": ProcessingFlow,
        "MAG": ProcessingFlow,
        "SSS": ProcessingFlow,
    }
    """
    result = {}
    for dt in ["SBP", "UHR", "MBES", "MAG", "SSS"]:
        flow = generate_flow_from_template(dt)
        flow.project_name = project_name
        flow.client = client
        flow.vessel = vessel
        flow.area = area
        result[dt] = flow
    return result


def generate_bulk_docx(flows, output_dir):
    """Generate DOCX reports for multiple flows.

    Input: {data_type: ProcessingFlow, ...}
    Writes files to output_dir as Processing_Report_{data_type}.docx
    Returns: list of output file paths.
    """
    os.makedirs(output_dir, exist_ok=True)
    output_paths = []
    for data_type, flow in flows.items():
        filename = f"Processing_Report_{data_type}.docx"
        output_path = os.path.join(output_dir, filename)
        generate_docx_report(flow, output_path)
        output_paths.append(output_path)
    return output_paths
